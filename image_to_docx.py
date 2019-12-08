from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_text('question number 1:\n')
r.add_picture('exemple.png')

document.save('exe.docx')



